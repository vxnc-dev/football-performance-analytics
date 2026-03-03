import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# CONFIG
CSV_PATH = r"" #SQUAD_ENRICHMENT.PY
OUTPUT_DOCX = r"" #PATH TO OUTPUT

df = pd.read_csv(CSV_PATH)

doc = Document()
doc.core_properties.author = "Weeno"
doc.core_properties.title = "Training Zones"

# TITLE
title = doc.add_paragraph()
run = title.add_run("Maximum Heart Rate and Heart Rate Zones")
run.bold = True
run.font.size = Pt(20)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)

# FORMULA
formula = doc.add_paragraph(
	"Sally-Edwards formula to determine maximum heart rate:\n"
	"214-(0.5 × Age)-(0.11 × Weight in kg)"
)
formula.paragraph_format.space_after = 18

position_order = ["Goalie", "Defender", "Midfield", "Attack"]

headers = [
	"Name",
	"Age",
	"HR-Max",
	"Zone 1\nLight zone",
	"Zone 2\nFat burning zone",
	"Zone 3\nAerobic zone",
	"Zone 4\nAnaerobic zone",
	"Zone 5\nMaximum Effort"
]

for position in position_order:
	group = df[df["position"] == position]
	if group.empty:
		continue

	# POSITION TITLE
	p = doc.add_paragraph()
	r = p.add_run(position)
	r.bold = True
	p.paragraph_format.space_before = 12
	p.paragraph_format.space_after = 6

	# TABLE
	table = doc.add_table(rows=1, cols=len(headers))
	table.style = "Table Grid"

	# HEADER ROW
	for i, h in enumerate(headers):
		cell = table.rows[0].cells[i]
		para = cell.paragraphs[0]
		para.text = h
		para.runs[0].bold = True
		para.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# DATA ROWS
	for _, row in group.iterrows():
		r = table.add_row().cells
		r[0].text = row["name"]
		r[1].text = str(row["age"])
		r[2].text = str(row["hr_max"])
		r[3].text = row["Zone 1 light zone"]
		r[4].text = row["Zone 2 fat burning zone"]
		r[5].text = row["Zone 3 aerobic zone"]
		r[6].text = row["Zone 4 anaerobic zone"]
		r[7].text = row["Zone 5 maximum effort"]

		for c in r:
			c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUTPUT_DOCX)
print("Word document created.")