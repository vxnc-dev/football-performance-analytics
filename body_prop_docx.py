from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# CONFIG
CSV_PATH = r"" #DATA_ENRICHMENT.PY
OUTPUT_DOCX = r"" #PATH TO OUTPUTS

# DATA
df = pd.read_csv(CSV_PATH)

mean_cols = ["age", "height_cm", "weight_kg", "bmi"]
squad_means = df[mean_cols].mean().round(2)
position_order = ["Goalie", "Defender", "Midfield", "Attack"]

doc = Document()
doc.core_properties.author = "Weeno"
doc.core_properties.title = "Body Proportions"

title = doc.add_paragraph()
run = title.add_run("Squad Body Proportions")
run.bold = True
run.font.size = Pt(20)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(18)

table = doc.add_table(rows=2, cols=4)
table.style = "Table Grid"

headers = ["Ø Age", "Ø Height", "Ø Weight", "Ø BMI"]
values = [
	squad_means["age"],
	squad_means["height_cm"],
	squad_means["weight_kg"],
	squad_means["bmi"],
]

for i, h in enumerate(headers):
	table.rows[0].cells[i].text = h
	table.rows[1].cells[i].text = str(values[i])

# HEADER
for cell in table.rows[0].cells:
	p = cell.paragraphs[0]
	p.runs[0].bold = True
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Value formatting
for cell in table.rows[1].cells:
	cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# POSITION SECTIONS
for position in position_order:
	group = df[df["position"] == position]
	if group.empty:
		continue

	# POSITION TITLE
	p = doc.add_paragraph()
	r = p.add_run(position)
	r.bold = True
	r.font.size = Pt(14)
	p.paragraph_format.space_before = Pt(14)
	p.paragraph_format.space_after = Pt(6)

	# PLAYER TABLE
	t = doc.add_table(rows=1, cols=8)
	t.style = "Table Grid"

	headers = ["Name", "Age", "Height", "Weight", "BMR", "Energy Requirememts (kcal)", "BMI", "BMI Category"]

	for i, h in enumerate(headers):
		cell = t.rows[0].cells[i]
		cell.text = h
		cp = cell.paragraphs[0]
		cp.runs[0].bold = True
		cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

	for _, row in group.iterrows():
		r = t.add_row().cells
		r[0].text = row["name"]
		r[1].text = str(row["age"])
		r[2].text = str(row["height_cm"])
		r[3].text = str(row["weight_kg"])
		r[4].text = str(row["bmr"])
		r[5].text = str(row["EBedarf"])
		r[6].text = str(row["bmi"])
		r[7].text = row["bmi_category"]

		for c in r:
			c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	doc.add_paragraph("")

doc.save(OUTPUT_DOCX)
print("Word document created successfully.")
